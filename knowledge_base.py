import os
import config_data as config
import hashlib
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from datetime import datetime


# ========== 文件文本提取函数 ==========
def extract_text_from_file(file_obj, file_ext: str = None) -> str:
    """从各种格式文件中提取文本

    Args:
        file_obj: 文件对象（可以是路径字符串或上传的文件对象）
        file_ext: 文件扩展名（如 'txt', 'pdf'），如果为 None 则从路径推断

    Returns:
        提取的文本内容
    """
    # 判断是文件路径还是文件对象
    if isinstance(file_obj, str):
        # 文件路径
        file_path = file_obj
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif ext == '.pdf':
                try:
                    import PyPDF2
                except ImportError:
                    raise ImportError("请安装 PyPDF2: pip install PyPDF2")

                text = ""
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                return text

            elif ext == '.docx':
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("请安装 python-docx: pip install python-docx")

                doc = Document(file_path)
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

            elif ext in ['.md', '.py']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            else:
                raise ValueError(f"不支持的文件格式: {ext}")

        except Exception as e:
            raise Exception(f"读取文件失败 {file_path}: {str(e)}")

    else:
        # 文件对象（如 Streamlit 上传的文件）
        if file_ext is None:
            raise ValueError("请提供文件扩展名")

        ext = file_ext.lower()

        try:
            if ext == 'txt':
                return file_obj.getvalue().decode('utf-8')

            elif ext == 'pdf':
                try:
                    import PyPDF2
                except ImportError:
                    raise ImportError("请安装 PyPDF2: pip install PyPDF2")

                text = ""
                reader = PyPDF2.PdfReader(file_obj)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                return text

            elif ext == 'docx':
                try:
                    from docx import Document
                except ImportError:
                    raise ImportError("请安装 python-docx: pip install python-docx")

                import io
                doc = Document(io.BytesIO(file_obj.getvalue()))
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

            elif ext in ['md', 'py']:
                return file_obj.getvalue().decode('utf-8')

            else:
                raise ValueError(f"不支持的文件格式: {ext}")

        except Exception as e:
            raise Exception(f"读取文件失败: {str(e)}")


# ========== MD5 相关函数 ==========
def check_md5(md5_str: str):
    """检查 MD5 是否已存在"""
    if not os.path.exists(config.md5_path):
        open(config.md5_path, 'w', encoding='utf-8').close()
        return False
    else:
        with open(config.md5_path, 'r', encoding='utf-8') as f:
            for line in f.readlines():
                line = line.strip()
                if line == md5_str:
                    return True
        return False


def save_md5(md5_str: str):
    """保存 MD5 到文件"""
    with open(config.md5_path, 'a', encoding='utf-8') as f:
        f.write(md5_str + '\n')


def get_string_md5(input_str: str, encoding='utf-8'):
    """计算字符串的 MD5"""
    str_bytes = input_str.encode(encoding=encoding)
    md5_obj = hashlib.md5()
    md5_obj.update(str_bytes)
    md5_hex = md5_obj.hexdigest()
    return md5_hex


# ========== 知识库服务类 ==========
class KnowledgeBaseService(object):
    def __init__(self):
        os.makedirs(config.persist_directory, exist_ok=True)

        # 使用配置中的 embedding 模型
        print(f"正在初始化 Embedding 模型: {config.embedding_model_name}")
        self.embedding = DashScopeEmbeddings(model=config.embedding_model_name)
        self.persist_directory = config.persist_directory
        self.vector_store = None
        self._load_or_create()

        # 文本分割器
        self.spliter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            separators=config.separators,
            length_function=len,
        )

    def _load_or_create(self):
        """加载或创建向量存储"""
        if os.path.exists(self.persist_directory) and os.listdir(self.persist_directory):
            try:
                self.vector_store = FAISS.load_local(
                    self.persist_directory,
                    self.embedding,
                    allow_dangerous_deserialization=True
                )
                print(f"✅ 已加载向量数据库，包含 {self.vector_store.index.ntotal} 个向量")
            except Exception as e:
                print(f"⚠️ 加载失败: {e}，创建新数据库")
                self.vector_store = FAISS.from_texts(["初始化知识库，请上传文档"], self.embedding)
                self.vector_store.save_local(self.persist_directory)
        else:
            print("📁 创建新的向量数据库")
            self.vector_store = FAISS.from_texts(["初始化知识库，请上传文档"], self.embedding)
            self.vector_store.save_local(self.persist_directory)

    def upload_by_str(self, data, filename):
        """上传文本到知识库"""
        if not data or not data.strip():
            return "[失败]内容为空，无法上传"

        md5_hex = get_string_md5(data)

        if check_md5(md5_hex):
            return "[跳过]内容已经存在知识库中"

        # 文本分割
        if len(data) > config.max_split_char_number:
            knowledge_chunks: list[str] = self.spliter.split_text(data)
        else:
            knowledge_chunks = [data]

        # 元数据
        metadata = {
            "source": filename,
            "create_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "operator": "小小",
            "chunk_count": len(knowledge_chunks)
        }

        # 添加到向量库
        self.vector_store.add_texts(
            knowledge_chunks,
            metadatas=[metadata for _ in knowledge_chunks],
        )
        self.vector_store.save_local(self.persist_directory)

        # 保存 MD5
        save_md5(md5_hex)

        return f"[成功]内容已经成功载入向量库，共 {len(knowledge_chunks)} 个片段"

    def get_retriever(self, k=None):
        """获取检索器，供 RAG 使用"""
        if k is None:
            k = config.similarity_threshold
        return self.vector_store.as_retriever(search_kwargs={"k": k})

    def get_vector_store(self):
        """获取向量存储实例"""
        return self.vector_store

    def get_stats(self):
        """获取知识库统计信息"""
        if self.vector_store and self.vector_store.index:
            return {
                "total_vectors": self.vector_store.index.ntotal,
                "persist_directory": self.persist_directory,
                "embedding_model": config.embedding_model_name
            }
        return None


if __name__ == '__main__':
    # 测试知识库服务
    service = KnowledgeBaseService()
    stats = service.get_stats()
    print(f"📊 知识库统计: {stats}")

    # 测试上传
    r = service.upload_by_str("您好，我是测试文本", "testfile.txt")
    print(r)

    # 测试检索
    retriever = service.get_retriever()
    results = retriever.invoke("测试")
    print(f"🔍 检索结果数: {len(results)}")